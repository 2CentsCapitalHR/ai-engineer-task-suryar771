import os
import json
from io import BytesIO
from typing import List, Dict, Any, Tuple

import streamlit as st
import streamlit.components.v1 as components
from dotenv import load_dotenv
load_dotenv(override=True)
import google.generativeai as genai

# Reuse the retrieval logic from query.py
from query import search_pinecone


def init_env() -> None:
    # Load .env once at startup and configure Gemini
    load_dotenv()
    gemini_key = os.getenv("GEMINI_API_KEY")
   # print(gemini_key)
    if gemini_key:
        genai.configure(api_key=gemini_key)


def render_match(idx: int, match: Dict[str, Any]) -> None:
    text = match.get("text") or match.get("metadata", {}).get("text") or ""
    score = match.get("score")
    meta = match.get("metadata", {})
    source = meta.get("source")
    url = meta.get("url")

    header_parts = [f"Chunk {idx}"]
    if source:
        header_parts.append(f"source: {source}")
    if score is not None:
        header_parts.append(f"score: {score:.4f}")
    with st.expander(" | ".join(header_parts), expanded=False):
        if url:
            st.write(f"URL: {url}")
        st.write(text)


def extract_docx_text_from_upload(upload) -> str:
    try:
        import docx  # python-docx
    except Exception as e:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx") from e

    # Read as in-memory bytes
    data = upload.read()
    document = docx.Document(BytesIO(data))
    parts: List[str] = []
    for p in document.paragraphs:
        if p.text:
            parts.append(p.text)
    return "\n".join(parts).strip()


def extract_docx_text_from_bytes(data: bytes) -> str:
    try:
        import docx  # python-docx
    except Exception as e:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx") from e

    document = docx.Document(BytesIO(data))
    parts: List[str] = []
    for p in document.paragraphs:
        if p.text:
            parts.append(p.text)
    return "\n".join(parts).strip()


def build_reviewed_docx_bytes(original_bytes: bytes, review_text: str) -> bytes:
    """Create an in-memory reviewed .docx by appending a review section."""
    try:
        import docx  # python-docx
    except Exception as e:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx") from e

    doc = docx.Document(BytesIO(original_bytes))
    doc.add_page_break()
    title_p = doc.add_paragraph()
    run = title_p.add_run("Review Comments (Automated)")
    run.bold = True
    for line in (review_text or "No comments.").splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.italic = True

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def build_review_prompt(context: str, doc_text: str) -> str:
    return (
        "You are an ADGM compliance assistant. Using ONLY the provided context (from the ADGM dataset), "
        "analyze the uploaded document for compliance. Follow the analytical style shown in Task.pdf.\n\n"
        "Produce two sections:\n"
        "1) Summary/Review: Paragraphs explaining compliance status, issues, and precise recommendations with ADGM law citations.\n"
        "2) Missing Documents: A bullet list of missing or incomplete mandatory documents/clauses.\n\n"
        "If the context does not support an answer, write: 'I could not find an answer in the provided documents.'\n\n"
        f"Context (retrieved from Pinecone):\n{context}\n\n"
        f"Company Submission (DOCX text):\n{doc_text}\n\n"
        "Respond in the exact two-section format described above."
    )


def review_docx_via_rag(doc_text: str) -> Tuple[str, List[Dict[str, Any]]]:
    # Build a retrieval query per requirement
    retrieval_query = (
        "Given this company submission text: "
        + doc_text[:8000] +
        ", identify missing mandatory documents and compliance issues."
    )

    matches = search_pinecone(retrieval_query, top_k=5)
    context = []
    for i, m in enumerate(matches):
        txt = (m.get("text") or m.get("metadata", {}).get("text") or "").strip()
        if not txt:
            continue
        context.append(f"[Chunk {i+1}]\n{txt}")
    context_text = "\n\n".join(context)

    # Use a lighter default to reduce quota pressure; allow override via .env
    model_name = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
    delay = 2
    prompt = build_review_prompt(context_text, doc_text[:12000])
    for attempt in range(4):
        try:
            model = genai.GenerativeModel(model_name)
            resp = model.generate_content(prompt)
            break
        except Exception as e:
            if attempt == 3:
                raise
            import time as _t
            _t.sleep(delay)
            delay = min(60, delay * 2)
    answer = getattr(resp, "text", None)
    if not answer and hasattr(resp, "candidates") and resp.candidates:
        try:
            parts = resp.candidates[0].content.parts
            answer = "".join(getattr(p, "text", "") for p in parts)
        except Exception:
            answer = None
    answer = (answer or "I could not find an answer in the provided documents.").strip()
    return answer, matches


def generate_structured_analysis(doc_text: str, matches: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Ask Gemini to return a strict JSON analysis for UI cards and lists."""
    context_blocks: List[str] = []
    for i, m in enumerate(matches):
        txt = (m.get("text") or m.get("metadata", {}).get("text") or "").strip()
        if txt:
            context_blocks.append(f"[Chunk {i+1}]\n{txt}")
    context_text = "\n\n".join(context_blocks)

    schema_hint = (
        "Return ONLY valid JSON (no markdown). Use this exact schema and keys:\n"
        "{\n"
        "  \"process\": \"Company Incorporation\",\n"
        "  \"documents_uploaded\": <int>,\n"
        "  \"required_documents\": <int>,\n"
        "  \"missing_documents\": [<string>],\n"
        "  \"issues_found\": [\n"
        "    {\n"
        "      \"document\": <string>,\n"
        "      \"section\": <string>,\n"
        "      \"issue\": <string>,\n"
        "      \"severity\": \"Critical|High|Medium|Low\",\n"
        "      \"suggestion\": <string>,\n"
        "      \"rule\": <string optional>\n"
        "    }\n"
        "  ]\n"
        "}\n"
    )

    prompt = (
        "You are an ADGM compliance assistant. From the retrieved reference context and the submitted document, "
        "produce a structured compliance analysis. Identify missing mandatory documents or clauses and list issues.\n\n"
        f"Context:\n{context_text[:12000]}\n\nDocument:\n{doc_text[:12000]}\n\n"
        + schema_hint
    )

    model_name = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
    delay = 2
    for attempt in range(4):
        try:
            model = genai.GenerativeModel(model_name)
            resp = model.generate_content(prompt)
            raw = (getattr(resp, "text", "") or "").strip()
            if not raw and hasattr(resp, "candidates") and resp.candidates:
                parts = resp.candidates[0].content.parts
                raw = "".join(getattr(p, "text", "") for p in parts)
            raw = raw.strip()
            if raw.startswith("```"):
                raw = raw.strip("`\n").lstrip("json").strip()
            data = json.loads(raw)
            # Basic normalization
            data.setdefault("process", "Company Incorporation")
            data.setdefault("documents_uploaded", 1)
            data.setdefault("required_documents", max(1, data.get("documents_uploaded", 1)))
            data.setdefault("missing_documents", [])
            data.setdefault("issues_found", [])
            return data
        except Exception:
            if attempt == 3:
                raise
            import time as _t
            _t.sleep(delay)
            delay = min(60, delay * 2)

    return {
        "process": "Company Incorporation",
        "documents_uploaded": 1,
        "required_documents": 1,
        "missing_documents": [],
        "issues_found": [],
    }


def compute_compliance_score(issues: List[Dict[str, Any]]) -> int:
    weights = {"critical": 25, "high": 15, "medium": 10, "low": 5}
    score = 100
    for issue in issues:
        sev = str(issue.get("severity", "")).lower()
        score -= weights.get(sev, 10)
    return max(0, min(100, score))




def main() -> None:
    init_env()
    st.set_page_config(page_title="ADGM Compliance Assistant", page_icon="📘", layout="centered")
    st.title("ADGM Compliance Assistant")
    st.subheader("Upload Corporate Documents")

    # Black & White theme with white accent
    st.markdown(
        """
        <style>
        :root {
          --bg:#000000; --card:#0a0a0a; --text:#ffffff; --muted:#b3b3b3; --border:rgba(255,255,255,0.2);
          --accent:#ffffff;
        }
        .block-container { padding-top: 1.25rem; }
        .stApp, body { background: var(--bg); color: var(--text); }
        /* Headings */
        h1, h2, h3, h4, h5, h6 { color: var(--text) !important; }
        /* File Uploader */
        .stFileUploader { border: 1px dashed var(--accent) !important; padding: 1.5rem !important; background: #050505; border-radius: 12px; }
        /* Button */
        .stButton>button { width: 100%; background: var(--accent); color:#000; font-weight:700; border-radius:8px; border:none; padding:0.6rem 1rem; }
        .stButton>button:hover { filter: none !important; }
        .stButton>button:disabled { opacity: 1 !important; pointer-events: none; cursor: default; filter: none !important; }
        /* Progress Bar (force white) */
        /* Track */
        .stProgress > div > div,
        div[data-testid="stProgressBar"] > div { background: rgba(255,255,255,0.15) !important; }
        /* Fill */
        .stProgress > div > div > div,
        div[data-testid="stProgressBar"] > div > div,
        div[role="progressbar"] {
          background: #ffffff !important;
          background-image: none !important;
        }
        /* Cards */
        .metric-card { background: var(--card); border: 1px solid var(--border); padding: 20px; border-radius: 10px; }
        .label { color: var(--accent); font-weight:600; font-size: 0.9rem; letter-spacing:0.02em; margin-bottom: 0; }
        .metric-card { transition: border-color .2s ease, transform .2s ease; min-height: 160px; display: flex; flex-direction: column; gap: 10px; }
        .metric-card:hover { border-color: rgba(255,255,255,0.35); transform: translateY(-1px); }
        .metric-value { font-size: 1.15rem; line-height: 1.6; color: var(--text); }
        .json-panel { background:#000; border:1px solid var(--border); border-radius:12px; padding:12px; position: relative; margin-top: -8px; }
        .json-toolbar { position:absolute; top:8px; right:8px; display:flex; gap:8px; }
        .btn-ghost { background: transparent; border:1px solid var(--accent); color: var(--accent); padding:6px 10px; border-radius:8px; cursor:pointer; }
        .btn-ghost:hover { background: rgba(255,255,255,0.08); }
        /* Ring */
        .ring { width:120px; height:120px; border-radius:50%; display:grid; place-items:center; margin:auto; border:6px solid var(--accent); }
        .ring .inner { background:#0a0a0a; width:88px; height:88px; border-radius:50%; display:grid; place-items:center; color:#fff; font-weight:700; }
        /* Pills */
        .pill { border:1px solid var(--accent); color:var(--accent); padding:2px 8px; border-radius:12px; font-size:12px; }
        /* Section hints */
        .hint { color: var(--muted); }
        /* Custom white loading bar directly under the Analyze button */
        .loader-bar { width: 100%; height: 5px; background: transparent; border-radius: 999px; overflow: hidden; margin-top: 8px; }
        .loader-bar .fill { height: 100%; width: 0%; background: #ffffff; transition: width 300ms ease; }
        /* Tabs styling */
        div[role="tablist"] { gap: 24px; border-bottom: 1px solid var(--border) !important; }
        button[role="tab"] { color: var(--muted) !important; background: transparent !important; border: none !important; }
        button[role="tab"][aria-selected="true"] { color: var(--accent) !important; }
        button[role="tab"][aria-selected="true"]::after { content: ""; display:block; height:3px; background: var(--accent); margin-top:8px; border-radius:2px; }
        /* Hide Streamlit header/toolbars (Deploy/nav) */
        div[data-testid="stToolbar"] { display: none !important; }
        header[data-testid="stHeader"] { display: none !important; }
        div[data-testid="stDecoration"] { display: none !important; }
        #MainMenu { visibility: hidden !important; }
        footer { visibility: hidden !important; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # File Upload Area
    st.markdown("<div class='label'>File Upload Area</div>", unsafe_allow_html=True)
    upload = st.file_uploader(" ", type=["docx"], accept_multiple_files=False)
    if "busy" not in st.session_state:
        st.session_state["busy"] = False
    if "analysis_running" not in st.session_state:
        st.session_state["analysis_running"] = False
    # Keep the button label constant; only disable while analysis is running
    st.session_state["busy"] = bool(st.session_state.get("analysis_running", False))
    go = st.button("Analyze for Compliance", use_container_width=True, disabled=st.session_state.get("analysis_running", False))
    loader_placeholder = st.empty()

    if go:
        # Immediately mark UI as busy and running
        st.session_state["busy"] = True
        st.session_state["analysis_running"] = True
        if not upload:
            st.warning("Please upload a .docx file.")
            st.stop()
        # Validate file type and size
        if not upload.name.lower().endswith(".docx"):
            st.error("Only .docx files are supported.")
            st.stop()
        # Read bytes once for reuse
        doc_bytes = upload.read()
        size_bytes = getattr(upload, "size", None) or len(doc_bytes)
        if size_bytes > 200 * 1024 * 1024:
            st.error("File exceeds 200MB limit.")
            st.stop()

        # Show white loader bar (initial)
        loader_placeholder.markdown("<div class='loader-bar'><div class='fill' style='width: 5%'></div></div>", unsafe_allow_html=True)
        try:
            doc_text = extract_docx_text_from_bytes(doc_bytes)
        except Exception as e:
            st.session_state["busy"] = False
            st.session_state["analysis_running"] = False
            st.error(f"Failed to read .docx: {e}")
            st.stop()

        if not doc_text.strip():
            st.warning("The uploaded document appears to be empty.")
            st.stop()

        # Update loader as retrieval begins
        loader_placeholder.markdown("<div class='loader-bar'><div class='fill' style='width: 30%'></div></div>", unsafe_allow_html=True)
        with st.spinner("Analyzing…"):
            try:
                review_answer, matches = review_docx_via_rag(doc_text)
            except Exception as e:
                st.session_state["busy"] = False
                st.session_state["analysis_running"] = False
                st.error(f"Error during review: {e}")
                st.stop()

        loader_placeholder.markdown("<div class='loader-bar'><div class='fill' style='width: 60%'></div></div>", unsafe_allow_html=True)
        # Structured analysis cards and lists
        try:
            analysis = generate_structured_analysis(doc_text, matches)
        except Exception:
            analysis = {
                "process": "Company Incorporation",
                "documents_uploaded": 1,
                "required_documents": 1,
                "missing_documents": [],
                "issues_found": [],
            }
        score = compute_compliance_score(analysis.get("issues_found", []))

        # Complete and hide loader
        loader_placeholder.markdown("<div class='loader-bar'><div class='fill' style='width: 100%'></div></div>", unsafe_allow_html=True)
        loader_placeholder.empty()

        # Build structured JSON output
        issues_src = analysis.get("issues_found") or []
        issues_min = [
            {
                "document": it.get("document"),
                "section": it.get("section"),
                "issue": it.get("issue"),
                "severity": it.get("severity"),
            }
            for it in issues_src
        ]
        structured_output = {
            "process": analysis.get("process", "Company Incorporation"),
            "documents_uploaded": analysis.get("documents_uploaded", 1),
            "required_documents": analysis.get("required_documents", 1),
            "missing_documents": analysis.get("missing_documents", []),
            "compliance_score": score,
            "issues_found": issues_min,
        }
        json_str = json.dumps(structured_output, ensure_ascii=False, indent=2)
        # STRUCTURED JSON OUTPUT panel with Copy + Download (reverted structure)
        st.markdown("<div class='label'>STRUCTURED JSON OUTPUT</div>", unsafe_allow_html=True)
       # st.markdown("<div class='json-panel' id='jsonPanel'>", unsafe_allow_html=True)
        components.html(
            f"""
            <div class='json-toolbar'>
              <button class='btn-ghost' id='copybtn'>Copy</button>
              <a id='downloadlink' class='btn-ghost' download='compliance_summary.json'>Download</a>
            </div>
            <script>
            const data = {json.dumps(json_str)};
            const blob = new Blob([data], {{ type: 'application/json' }});
            const url = URL.createObjectURL(blob);
            const dl = document.getElementById('downloadlink');
            if (dl) dl.href = url;
            const btn = document.getElementById('copybtn');
            if (btn) {{ btn.onclick = async () => {{ try {{ await navigator.clipboard.writeText(data); btn.innerText='Copied'; }} catch(e) {{ btn.innerText='Copy Failed'; }} }} }}
            </script>
            """,
            height=0,
        )
        st.code(json_str, language="json")
        st.markdown("</div>", unsafe_allow_html=True)

        # 3-column metrics row: Process | Documents | Compliance Score
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown(
                "<div class='metric-card'><div class='label'>Process</div>" +
                f"<div class='metric-value' style='margin-top:6px; font-size:1.8rem; font-weight:700'>{analysis.get('process','Company Incorporation')}</div></div>",
                unsafe_allow_html=True,
            )
        with c2:
            st.markdown(
                "<div class='metric-card'><div class='label'>Documents</div>" +
                f"<div class='metric-value' style='margin-top:6px; font-size:2.2rem; font-weight:800'>{analysis.get('documents_uploaded',1)} / {analysis.get('required_documents',1)}</div></div>",
                unsafe_allow_html=True,
            )
        with c3:
            pct = max(0, min(100, score))
            st.markdown(
                f"<div class='metric-card'><div class='label'>Compliance Score</div>"
                f"<div class='metric-value' style='margin-top:10px; font-size:2.2rem; font-weight:800'>{pct}/100</div></div>",
                unsafe_allow_html=True,
            )

        # Tabs for results
        tabs = st.tabs(["Analysis Summary", "Missing Documents", "Issues Found"])

        # Analysis Summary tab with downloads
        with tabs[0]:
            try:
                reviewed_bytes = build_reviewed_docx_bytes(doc_bytes, review_answer)
                st.download_button(
                    "Download Reviewed DOCX",
                    data=reviewed_bytes,
                    file_name=f"reviewed_{upload.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception:
                pass
            st.download_button("Download JSON Summary", data=json_str, file_name="compliance_summary.json", mime="application/json")

        # Missing Documents tab
        with tabs[1]:
            missing_docs = analysis.get("missing_documents") or []
            if isinstance(missing_docs, list) and missing_docs:
                for md in missing_docs:
                    st.markdown(f"- {md}")
            else:
                st.caption("No missing documents detected.")

        # Issues Found tab
        with tabs[2]:
            issues = analysis.get("issues_found") or []
            if not issues:
                st.caption("No issues detected.")
            for issue in issues:
                doc = issue.get("document") or "Document"
                section = issue.get("section") or ""
                title = f"{doc} - {section}".strip(" -")
                sev = (issue.get("severity") or "").upper()
                sev_color = "#ffffff"
                st.markdown(
                    f"<div style='border:1px solid rgba(255,255,255,0.15); background:#0a0a0a; padding:12px; border-radius:8px; margin-bottom:10px;'>"
                    f"<div style='display:flex; justify-content:space-between; align-items:center;'>"
                    f"<strong>{title}</strong>"
                    f"<span class='pill' style='border-color:{sev_color}; color:{sev_color};'>{sev or 'N/A'}</span>"
                    f"</div>",
                    unsafe_allow_html=True,
                )
                if issue.get("issue"):
                    st.markdown(f"<div style='margin-top:6px;'>{issue['issue']}</div>", unsafe_allow_html=True)
                if issue.get("suggestion"):
                    st.markdown(f"<div style='margin-top:6px;'><em>Suggestion:</em> {issue['suggestion']}</div>", unsafe_allow_html=True)
                if issue.get("rule"):
                    st.markdown(f"<div style='margin-top:6px;'><strong>Rule:</strong> {issue['rule']}</div>", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

        # Clear busy state after finishing render
        st.session_state["busy"] = False
        st.session_state["analysis_running"] = False



if __name__ == "__main__":
    main()

