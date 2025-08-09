import os
import sys
import json
import time
from typing import List, Dict, Any, Tuple

from dotenv import load_dotenv
from pinecone import Pinecone
import google.generativeai as genai

try:
    import docx  # python-docx
except Exception:
    docx = None  # type: ignore


INDEX_NAME = "adgm-compliance"  # Pinecone index (dim=768, cosine)
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.0-pro")
GEMINI_EMBED_MODEL = "text-embedding-004"  # 768-dim
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reviewed_docs")


def log(message: str) -> None:
    ts = time.strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{ts}] {message}")


def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def init_clients() -> Pinecone:
    load_dotenv()
    gemini_key = os.getenv("GEMINI_API_KEY")
    pinecone_key = os.getenv("PINECONE_API_KEY")
    if not gemini_key:
        raise RuntimeError("Missing GEMINI_API_KEY in .env")
    if not pinecone_key:
        raise RuntimeError("Missing PINECONE_API_KEY in .env")
    genai.configure(api_key=gemini_key)
    return Pinecone(api_key=pinecone_key)


def get_index(pc: Pinecone):
    idx = pc.Index(INDEX_NAME)
    try:
        desc = pc.describe_index(INDEX_NAME)
        dim = (
            (desc.get("dimension") if isinstance(desc, dict) else None)
            or (desc.get("spec", {}).get("pod", {}).get("dimension") if isinstance(desc, dict) else None)
            or 768
        )
        if int(dim) != 768:
            raise RuntimeError(f"Index '{INDEX_NAME}' has dimension {dim}, expected 768 for Gemini embeddings")
    except Exception:
        # Proceed; assume configured correctly
        pass
    return idx


def embed_text(text: str) -> List[float]:
    # For safety, truncate excessively long texts
    if len(text) > 200_000:
        text = text[:200_000]
    res = genai.embed_content(model=GEMINI_EMBED_MODEL, content=text)
    emb = res.get("embedding") if isinstance(res, dict) else getattr(res, "embedding", None)
    if isinstance(emb, dict):
        emb = emb.get("values")
    if emb is None:
        raise RuntimeError("Failed to create Gemini embedding")
    return list(emb)


def pinecone_search(index, text: str, top_k: int = 5) -> List[Dict[str, Any]]:
    vec = embed_text(text)
    res = index.query(vector=vec, top_k=top_k, include_values=False, include_metadata=True)
    matches = []
    for m in (res.get("matches", []) if isinstance(res, dict) else getattr(res, "matches", [])):
        md = m.get("metadata", {}) if isinstance(m, dict) else getattr(m, "metadata", {})
        score = m.get("score") if isinstance(m, dict) else getattr(m, "score", None)
        text_chunk = md.get("text") or ""
        matches.append({"text": text_chunk, "score": score, "metadata": md})
    return matches


def extract_docx_text(path: str) -> str:
    if docx is None:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")
    document = docx.Document(path)
    parts: List[str] = []
    for p in document.paragraphs:
        if p.text:
            parts.append(p.text)
    return "\n".join(parts).strip()


def classify_doc_type(text: str) -> str:
    lt = text.lower()
    # Simple keyword rules
    if "memorandum of association" in lt or "moa" in lt:
        return "Memorandum of Association"
    if "articles of association" in lt or "aoa" in lt:
        return "Articles of Association"
    if "board resolution" in lt or "resolution of the board" in lt:
        return "Board Resolution"

    # Fallback to Gemini classification
    prompt = (
        "You will classify the type of an ADGM corporate document. Choose exactly one label from: "
        "['Memorandum of Association', 'Articles of Association', 'Board Resolution', 'Other']. "
        "Respond with the label only, no extra words.\n\n"
        f"Document:\n{text[:6000]}\n"
    )
    model = genai.GenerativeModel(GEMINI_MODEL)
    resp = model.generate_content(prompt)
    label = (getattr(resp, "text", "") or "").strip()
    label = label.strip('"').strip("'")
    valid = {"Memorandum of Association", "Articles of Association", "Board Resolution", "Other"}
    if label not in valid:
        return "Other"
    return label


def load_company_incorporation_checklist(index) -> List[str]:
    # Retrieve relevant chunks
    query = (
        "ADGM company incorporation required document checklist: list the required documents "
        "for incorporation, such as Memorandum of Association, Articles of Association, Board Resolution, etc."
    )
    matches = pinecone_search(index, query, top_k=10)
    context = "\n\n".join(m["text"] for m in matches if m.get("text"))
    if not context:
        # Fallback static minimal checklist
        return [
            "Memorandum of Association",
            "Articles of Association",
            "Board Resolution",
        ]

    prompt = (
        "From the context below, extract the checklist of documents typically required for ADGM Company Incorporation. "
        "Return a simple bullet list with one document name per line, without explanations.\n\n"
        f"Context:\n{context[:12000]}\n"
    )
    model = genai.GenerativeModel(GEMINI_MODEL)
    resp = model.generate_content(prompt)
    text = (getattr(resp, "text", "") or "").strip()
    items: List[str] = []
    for line in text.splitlines():
        line = line.strip("-• \t").strip()
        if not line:
            continue
        # Normalize capitalization
        if len(line) > 200:
            continue
        items.append(line)

    # Deduplicate and normalize common names
    norm_map = {
        "moa": "Memorandum of Association",
        "memorandum of association": "Memorandum of Association",
        "articles of association": "Articles of Association",
        "aoa": "Articles of Association",
        "board resolution": "Board Resolution",
    }
    normalized: List[str] = []
    seen = set()
    for it in items:
        key = it.lower()
        name = norm_map.get(key, it)
        if name not in seen:
            seen.add(name)
            normalized.append(name)
    if not normalized:
        normalized = [
            "Memorandum of Association",
            "Articles of Association",
            "Board Resolution",
        ]
    return normalized


def check_missing_docs(uploaded_doc_types: List[str], checklist: List[str]) -> List[str]:
    norm = {
        "memorandum of association": "Memorandum of Association",
        "articles of association": "Articles of Association",
        "board resolution": "Board Resolution",
    }
    uploaded_norm = {norm.get(d.lower(), d) for d in uploaded_doc_types}
    missing = []
    for item in checklist:
        canon = norm.get(item.lower(), item)
        if canon not in uploaded_norm:
            missing.append(canon)
    return missing


def rag_review_doc(doc_text: str, index) -> str:
    # Retrieve top reference chunks given the document content
    matches = pinecone_search(index, doc_text, top_k=5)
    context = "\n\n".join(m["text"] for m in matches if m.get("text"))
    if not context:
        context = ""

    prompt = (
        "Review this document for ADGM compliance based on retrieved reference text. "
        "Flag missing clauses, legal issues, and suggest specific changes with ADGM law citations.\n\n"
        f"Retrieved Reference Text:\n{context[:12000]}\n\n"
        f"Document To Review:\n{doc_text[:12000]}\n\n"
        "Provide a concise list of findings and recommended edits."
    )
    model = genai.GenerativeModel(GEMINI_MODEL)
    resp = model.generate_content(prompt)
    review = (getattr(resp, "text", "") or "").strip()
    return review or "No issues found."


def add_comments_to_docx(input_path: str, comments_text: str, output_path: str) -> None:
    if docx is None:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")
    document = docx.Document(input_path)

    # Append a review section at the end (inline-style simulated comments)
    document.add_page_break()
    heading = document.add_paragraph()
    run = heading.add_run("Review Comments (Automated)")
    run.bold = True

    # Split comments into lines/points
    lines = [ln.strip() for ln in comments_text.splitlines() if ln.strip()]
    if not lines:
        lines = ["No issues found."]

    for ln in lines:
        p = document.add_paragraph()
        r = p.add_run(ln)
        r.italic = True

    ensure_dir(os.path.dirname(output_path))
    document.save(output_path)


def main() -> None:
    if len(sys.argv) < 2:
        print("Usage: python compliance_check.py <file1.docx> [file2.docx ...]")
        sys.exit(1)

    pc = init_clients()
    index = get_index(pc)
    ensure_dir(OUTPUT_DIR)

    uploaded_docs: List[Dict[str, str]] = []
    issues_found: List[str] = []

    for path in sys.argv[1:]:
        if not os.path.isfile(path) or not path.lower().endswith(".docx"):
            log(f"SKIP: Not a .docx file: {path}")
            continue

        filename = os.path.basename(path)
        log(f"Processing: {filename}")

        try:
            text = extract_docx_text(path)
        except Exception as e:
            log(f"ERROR: Failed to read {filename}: {e}")
            continue

        doc_type = classify_doc_type(text)
        uploaded_docs.append({"file": filename, "doc_type": doc_type})
        log(f"Classified '{filename}' as: {doc_type}")

        try:
            review = rag_review_doc(text, index)
            issues_found.append(f"{filename}: {review}")
        except Exception as e:
            log(f"ERROR: RAG review failed for {filename}: {e}")
            review = f"Review failed: {e}"

        reviewed_name = os.path.splitext(filename)[0] + "_reviewed.docx"
        reviewed_path = os.path.join(OUTPUT_DIR, reviewed_name)
        try:
            add_comments_to_docx(path, review, reviewed_path)
            log(f"Saved reviewed document: {reviewed_path}")
        except Exception as e:
            log(f"ERROR: Failed to write reviewed DOCX for {filename}: {e}")

    # Build checklist and detect missing docs
    try:
        checklist = load_company_incorporation_checklist(index)
        missing_docs = check_missing_docs([d["doc_type"] for d in uploaded_docs], checklist)
    except Exception as e:
        log(f"ERROR: Checklist evaluation failed: {e}")
        checklist = []
        missing_docs = []

    summary = {
        "process": "Company Incorporation",
        "uploaded_docs": uploaded_docs,
        "missing_docs": missing_docs,
        "issues_found": issues_found,
    }

    summary_path = os.path.join(OUTPUT_DIR, "compliance_summary.json")
    try:
        with open(summary_path, "w", encoding="utf-8") as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)
        log(f"Saved summary: {summary_path}")
    except Exception as e:
        log(f"ERROR: Failed to write summary JSON: {e}")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        log("Interrupted by user.")
        try:
            sys.exit(1)
        except SystemExit:
            os._exit(1)

